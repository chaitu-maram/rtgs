from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create new document
doc = Document()

# Title
title = doc.add_paragraph("AXIS BANK RTGS / NEFT / IMPS REQUEST FORM")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].bold = True
title.runs[0].font.size = Pt(14)

# Branch and date
doc.add_paragraph("\nTo\t\tDate: {{date}}")
doc.add_paragraph("The Branch Head")
doc.add_paragraph("{{branch}} Branch")
doc.add_paragraph("Dear Sir/Madam,")

# PAN and amount
doc.add_paragraph("\nPAN No.: {{remitter_pan}}")
doc.add_paragraph("Please remit through RTGS / NEFT / IMPS a sum of Rs. {{amount}} /- ({{amount_words}} only), as per the details given below:")

# Remitter section
doc.add_paragraph("\nRemitter Details").runs[0].bold = True
doc.add_paragraph("Name: {{remitter_name}}")
doc.add_paragraph("Address: {{remitter_address}}")
doc.add_paragraph("Mobile Number: {{remitter_mobile}}")
doc.add_paragraph("Account Number: {{remitter_account}}")
doc.add_paragraph("Cheque Number: {{cheque_number}}")
doc.add_paragraph("Cheque Date: {{cheque_date}}")

# Beneficiary section
doc.add_paragraph("\nBeneficiary Details").runs[0].bold = True
doc.add_paragraph("Name: {{beneficiary_name}}")
doc.add_paragraph("Account Number: {{beneficiary_account}}")
doc.add_paragraph("Bank Name: {{beneficiary_bank}}")
doc.add_paragraph("Branch Name: {{beneficiary_branch}}")
doc.add_paragraph("IFSC Code: {{beneficiary_ifsc}}")
doc.add_paragraph("Beneficiary Type: {{beneficiary_type}}")
doc.add_paragraph("LEI Code (if applicable): {{beneficiary_lei}}")
doc.add_paragraph("Sender to Receiver Information: {{sender_to_receiver}}")

# Signature block
doc.add_paragraph("\nI / We hereby authorize Axis Bank Ltd to carry out RTGS/NEFT/IMPS remittance as per the details mentioned above and to debit my/our account for the amount plus applicable charges and taxes.")
doc.add_paragraph("\nCustomer Signature(s): ___________________________")
doc.add_paragraph("Authorised Signatory (Affix Stamp in case of Non-individual): ___________________________")

# Save to file
doc.save("Standard_RTGS_Template.docx")
